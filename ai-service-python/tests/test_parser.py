"""
Tests for DocumentParser

Covers all supported formats: PDF, DOCX, XLSX, TXT, MD.
Also tests error cases: unsupported format, empty file, corrupted file.

Run::

    pytest tests/test_parser.py -v
    python -m pytest tests/test_parser.py -v
"""

from __future__ import annotations

import io
import tempfile
from pathlib import Path

import pytest

from app.services.parser import (
    DocumentParser,
    EmptyDocumentError,
    ParserError,
    UnsupportedFormatError,
    parse_document,
)


# ============================================================
# Fixtures — generate test files in memory
# ============================================================

@pytest.fixture
def parser() -> DocumentParser:
    """Return a fresh DocumentParser instance."""
    return DocumentParser()


# ---- PDF ----

@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Generate a multi-page PDF with real text using fpdf2.

    Uses ASCII content because built-in fonts don't support CJK.
    Chinese text extraction is covered by TXT/MD tests.
    """
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Page 1 — has text
    pdf.add_page()
    pdf.set_font("Helvetica", "B", size=14)
    pdf.cell(0, 10, "Chapter 1: Project Overview", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, "This is a Knowledge RAG system for enterprise knowledge management.")

    # Page 2 — has text
    pdf.add_page()
    pdf.set_font("Helvetica", "B", size=14)
    pdf.cell(0, 10, "Chapter 2: Technical Architecture", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, "Backend: Spring Boot + MyBatis-Plus. AI Service: FastAPI + Chroma.")

    # Page 3 — empty (will be filtered)
    pdf.add_page()

    # Page 4 — has text
    pdf.add_page()
    pdf.set_font("Helvetica", "B", size=14)
    pdf.cell(0, 10, "Chapter 4: Deployment", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, "Deploy with Docker Compose. One-command startup.")

    return pdf.output()


# ---- DOCX ----

@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Generate a simple DOCX with paragraphs and a table."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("产品需求文档", level=1)
    doc.add_paragraph("版本：v2.0")
    doc.add_paragraph("本文档描述了知识库RAG系统的核心需求。")
    doc.add_heading("功能需求", level=2)
    doc.add_paragraph("1. 知识库CRUD —— 支持创建、查询、更新、软删除。")
    doc.add_paragraph("2. 文档上传 —— 支持 PDF/DOCX/XLSX/TXT/MD 格式。")
    doc.add_paragraph("3. 智能问答 —— 基于RAG的语义检索和LLM生成。")

    # Add a small table
    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["模块", "技术栈", "状态"]
    data = [
        ["知识库管理", "Spring Boot + MyBatis-Plus", "已完成"],
        ["文档解析", "FastAPI + pypdf", "开发中"],
    ]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---- XLSX ----

@pytest.fixture
def sample_xlsx_bytes() -> bytes:
    """Generate a simple XLSX with two sheets."""
    from openpyxl import Workbook

    wb = Workbook()

    # Sheet 1 — 项目信息
    ws1 = wb.active
    ws1.title = "项目信息"
    ws1.append(["项目名称", "版本", "负责人"])
    ws1.append(["知识库RAG系统", "v1.0", "张三"])
    ws1.append(["AI服务", "v1.0", "李四"])

    # Sheet 2 — 技术指标
    ws2 = wb.create_sheet("技术指标")
    ws2.append(["指标", "目标值", "当前值"])
    ws2.append(["检索延迟", "<500ms", "320ms"])
    ws2.append(["准确率", ">90%", "93.5%"])
    ws2.append(["并发用户", "100", "已达标"])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---- TXT ----

@pytest.fixture
def sample_txt_bytes() -> bytes:
    """Generate a UTF-8 TXT file with Chinese content."""
    text = (
        "企业知识库管理系统\n"
        "==================\n\n"
        "1. 概述\n"
        "   本系统旨在构建企业级知识管理平台，\n"
        "   支持文档存储、智能检索和问答。\n\n"
        "2. 核心功能\n"
        "   - 知识库 CRUD\n"
        "   - 文档上传与解析\n"
        "   - RAG 智能问答\n"
    )
    return text.encode("utf-8")


# ---- MD ----

@pytest.fixture
def sample_md_bytes() -> bytes:
    """Generate a Markdown file with Chinese content."""
    text = (
        "# 技术文档\n\n"
        "## 架构设计\n\n"
        "系统采用**微服务架构**，主要包含以下模块：\n\n"
        "- **Java 后端**: Spring Boot + MyBatis-Plus\n"
        "- **Python AI 服务**: FastAPI + Chroma\n"
        "- **存储层**: MySQL + MinIO + Redis\n\n"
        "```python\n"
        "print('Hello, RAG!')\n"
        "```\n\n"
        "> 注意：本文档为技术概要，详细设计见各模块文档。\n"
    )
    return text.encode("utf-8")


# ---- GBK-encoded TXT ----

@pytest.fixture
def sample_gbk_bytes() -> bytes:
    """Generate a GBK-encoded TXT file (common for Chinese documents)."""
    text = "这是一份GBK编码的中文文档。\n用于测试编码自动检测。\n"
    return text.encode("gbk")


# ============================================================
# Tests — Happy Path
# ============================================================

class TestPdfParsing:
    """PDF parsing tests."""

    def test_parse_pdf_basic(self, parser, sample_pdf_bytes):
        text = parser.parse(sample_pdf_bytes, "pdf")
        assert "Chapter 1" in text
        assert "Chapter 2" in text
        assert "Chapter 4" in text

    def test_parse_pdf_filters_empty_pages(self, parser, sample_pdf_bytes):
        text = parser.parse(sample_pdf_bytes, "pdf")
        # Page 3 was empty — should NOT appear
        assert "Chapter 3" not in text

    def test_parse_pdf_case_insensitive_type(self, parser, sample_pdf_bytes):
        text = parser.parse(sample_pdf_bytes, "PDF")
        assert "Chapter 1" in text

    def test_parse_pdf_via_convenience_function(self, sample_pdf_bytes):
        text = parse_document(sample_pdf_bytes, "pdf")
        assert len(text) > 0


class TestDocxParsing:
    """DOCX parsing tests."""

    def test_parse_docx_paragraphs(self, parser, sample_docx_bytes):
        text = parser.parse(sample_docx_bytes, "docx")
        assert "产品需求文档" in text
        assert "知识库CRUD" in text
        assert "智能问答" in text

    def test_parse_docx_table_content(self, parser, sample_docx_bytes):
        text = parser.parse(sample_docx_bytes, "docx")
        assert "知识库管理" in text
        assert "Spring Boot" in text
        assert "已完成" in text


class TestXlsxParsing:
    """XLSX parsing tests."""

    def test_parse_xlsx_sheet_headings(self, parser, sample_xlsx_bytes):
        text = parser.parse(sample_xlsx_bytes, "xlsx")
        assert "## Sheet: 项目信息" in text
        assert "## Sheet: 技术指标" in text

    def test_parse_xlsx_data_rows(self, parser, sample_xlsx_bytes):
        text = parser.parse(sample_xlsx_bytes, "xlsx")
        assert "张三" in text
        assert "93.5%" in text
        assert "已达标" in text

    def test_parse_xlsx_tab_separated(self, parser, sample_xlsx_bytes):
        text = parser.parse(sample_xlsx_bytes, "xlsx")
        # Rows use tab as separator
        assert "\t" in text


class TestPlainTextParsing:
    """TXT and MD parsing tests."""

    def test_parse_txt_utf8(self, parser, sample_txt_bytes):
        text = parser.parse(sample_txt_bytes, "txt")
        assert "企业知识库管理系统" in text
        assert "RAG 智能问答" in text

    def test_parse_md_content(self, parser, sample_md_bytes):
        text = parser.parse(sample_md_bytes, "md")
        assert "# 技术文档" in text
        assert "Spring Boot" in text
        assert "```python" in text

    def test_parse_txt_gbk_fallback(self, parser, sample_gbk_bytes):
        text = parser.parse(sample_gbk_bytes, "txt")
        assert "GBK编码" in text
        assert "编码自动检测" in text


# ============================================================
# Tests — Error Cases
# ============================================================

class TestErrorHandling:
    """Error handling tests."""

    def test_unsupported_format(self, parser):
        with pytest.raises(UnsupportedFormatError) as exc:
            parser.parse(b"dummy content", "pptx")
        assert "不支持的文件格式" in str(exc.value)
        assert "pptx" in str(exc.value)

    def test_unsupported_format_unknown(self, parser):
        with pytest.raises(UnsupportedFormatError) as exc:
            parser.parse(b"content", "html")
        assert "html" in str(exc.value)

    def test_empty_bytes(self, parser):
        with pytest.raises(ParserError) as exc:
            parser.parse(b"", "pdf")
        assert "文件内容为空" in str(exc.value)

    def test_corrupted_pdf(self, parser):
        with pytest.raises(ParserError) as exc:
            parser.parse(b"this is not a pdf file", "pdf")
        assert "pdf" in str(exc.value).lower()

    def test_corrupted_docx(self, parser):
        with pytest.raises(ParserError) as exc:
            parser.parse(b"not a valid docx zip", "docx")
        assert "docx" in str(exc.value).lower()

    def test_corrupted_xlsx(self, parser):
        with pytest.raises(ParserError) as exc:
            parser.parse(b"not a valid xlsx zip", "xlsx")
        assert "xlsx" in str(exc.value).lower()

    def test_empty_pdf_no_text(self, parser):
        """PDF with pages but no text raises EmptyDocumentError."""
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.add_page()
        empty_pdf = pdf.output()

        with pytest.raises(EmptyDocumentError) as exc:
            parser.parse(empty_pdf, "pdf")
        assert "未提取到任何文本" in str(exc.value)


class TestPostProcessing:
    """Post-processing tests."""

    def test_collapse_excessive_newlines(self, parser):
        text = "第一段\n\n\n\n第二段\n\n\n\n\n第三段"
        result = parser._postprocess(text)
        # 3+ newlines → 2
        assert "\n\n\n" not in result
        assert result == "第一段\n\n第二段\n\n第三段"

    def test_strip_whitespace(self, parser):
        text = "\n\n  内容  \n\n"
        result = parser._postprocess(text)
        assert result == "内容"


# ============================================================
# Tests — Integration-style (with temp files)
# ============================================================

class TestWithTempFiles:
    """Tests that write bytes to disk and parse by re-reading."""

    def test_roundtrip_txt(self, parser, sample_txt_bytes, tmp_path):
        """Write TXT to disk, read back, parse."""
        file_path = tmp_path / "test.txt"
        file_path.write_bytes(sample_txt_bytes)

        read_back = file_path.read_bytes()
        text = parser.parse(read_back, "txt")
        assert "企业知识库管理系统" in text

    def test_roundtrip_md(self, parser, sample_md_bytes, tmp_path):
        file_path = tmp_path / "README.md"
        file_path.write_bytes(sample_md_bytes)

        read_back = file_path.read_bytes()
        text = parser.parse(read_back, "md")
        assert "架构设计" in text

    def test_roundtrip_pdf(self, parser, sample_pdf_bytes, tmp_path):
        file_path = tmp_path / "doc.pdf"
        file_path.write_bytes(sample_pdf_bytes)

        read_back = file_path.read_bytes()
        text = parser.parse(read_back, "pdf")
        assert "Chapter 1" in text

    def test_roundtrip_docx(self, parser, sample_docx_bytes, tmp_path):
        file_path = tmp_path / "doc.docx"
        file_path.write_bytes(sample_docx_bytes)

        read_back = file_path.read_bytes()
        text = parser.parse(read_back, "docx")
        assert "产品需求文档" in text

    def test_roundtrip_xlsx(self, parser, sample_xlsx_bytes, tmp_path):
        file_path = tmp_path / "data.xlsx"
        file_path.write_bytes(sample_xlsx_bytes)

        read_back = file_path.read_bytes()
        text = parser.parse(read_back, "xlsx")
        assert "## Sheet: 项目信息" in text


# ============================================================
# Tests — Edge Cases
# ============================================================

class TestEdgeCases:
    """Edge case tests."""

    def test_large_text(self, parser):
        """Parser should handle large text without issues."""
        large = "这是测试文本。\n" * 10000
        data = large.encode("utf-8")
        text = parser.parse(data, "txt")
        assert len(text) > 50000

    def test_text_with_special_characters(self, parser):
        special = "特殊字符: © ® ™ — … 「」 『』 【】"
        data = special.encode("utf-8")
        text = parser.parse(data, "txt")
        assert "©" in text
        assert "【】" in text

    def test_text_with_numbers_and_urls(self, parser):
        content = (
            "参考链接：\n"
            "https://example.com/doc/v1.0\n"
            "联系电话：+86-10-1234-5678\n"
            "版本号：v3.14.159"
        )
        data = content.encode("utf-8")
        text = parser.parse(data, "txt")
        assert "https://example.com" in text
        assert "v3.14.159" in text
