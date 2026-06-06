"""
Document Parser Service

Extracts plain text from uploaded documents.
Supported formats: PDF, DOCX, XLSX, TXT, MD.

Delegates to format-specific libraries:
- PDF  → pypdf
- DOCX → python-docx
- XLSX → openpyxl
- TXT  → built-in (with encoding detection fallback)
- MD   → built-in (same as TXT)

Design note:
    This module returns raw extracted text only.
    Chunking and embedding are handled in later chapters.
"""

from __future__ import annotations

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


# ============================================================
# Exception Classes
# ============================================================

class UnsupportedFormatError(ValueError):
    """Raised when the file type is not in the supported list."""

    def __init__(self, file_type: str) -> None:
        self.file_type = file_type
        super().__init__(
            f"不支持的文件格式: '{file_type}'。"
            f"支持的类型: PDF, DOCX, XLSX, TXT, MD"
        )


class ParserError(RuntimeError):
    """Raised when text extraction fails for a supported format."""

    def __init__(self, file_type: str, detail: str) -> None:
        self.file_type = file_type
        self.detail = detail
        super().__init__(f"文档解析失败 [{file_type}]: {detail}")


class EmptyDocumentError(ParserError):
    """Raised when a document contains no extractable text."""

    def __init__(self, file_type: str) -> None:
        super().__init__(file_type, "文档中未提取到任何文本内容")


# ============================================================
# DocumentParser
# ============================================================

class DocumentParser:
    """
    Multi-format document text extractor.

    Usage::

        parser = DocumentParser()
        text = parser.parse(file_bytes, "pdf")

    All ``parse_*`` methods return plain text as a single string.
    """

    # Supported file types
    SUPPORTED_TYPES = frozenset({"pdf", "docx", "xlsx", "txt", "md"})

    # ---- Public API ----

    def parse(self, file_bytes: bytes, file_type: str) -> str:
        """
        Extract text from file bytes.

        Args:
            file_bytes: Raw file content.
            file_type:  Lowercase extension, e.g. "pdf", "docx".

        Returns:
            Extracted plain text.

        Raises:
            UnsupportedFormatError: file_type not supported.
            ParserError:           extraction failed.
            EmptyDocumentError:    no text found.
        """
        file_type = file_type.lower().strip()

        if file_type not in self.SUPPORTED_TYPES:
            raise UnsupportedFormatError(file_type)

        if not file_bytes:
            raise ParserError(file_type, "文件内容为空")

        logger.info("开始解析文档: type=%s, size=%d bytes", file_type, len(file_bytes))

        # Dispatch to format-specific parser
        try:
            if file_type == "pdf":
                text = self._parse_pdf(file_bytes)
            elif file_type == "docx":
                text = self._parse_docx(file_bytes)
            elif file_type == "xlsx":
                text = self._parse_xlsx(file_bytes)
            elif file_type in ("txt", "md"):
                text = self._parse_plain_text(file_bytes, file_type)
            else:
                raise UnsupportedFormatError(file_type)  # defensive
        except (UnsupportedFormatError, ParserError, EmptyDocumentError):
            raise
        except Exception as e:
            logger.exception("解析异常: type=%s", file_type)
            raise ParserError(file_type, str(e)) from e

        # Post-processing
        text = self._postprocess(text)

        if not text.strip():
            raise EmptyDocumentError(file_type)

        logger.info("文档解析完成: type=%s, text_length=%d", file_type, len(text))
        return text

    # ---- PDF Parser ----

    def _parse_pdf(self, file_bytes: bytes) -> str:
        """
        Extract text from PDF using pypdf.

        Filters out empty pages (pages whose extracted text,
        after stripping, is empty).
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ParserError("pdf", "pypdf 库未安装，请执行: pip install pypdf")

        stream = io.BytesIO(file_bytes)
        try:
            reader = PdfReader(stream)
        except Exception as e:
            raise ParserError("pdf", f"无法读取 PDF 文件（文件可能已损坏）: {e}")

        total_pages = len(reader.pages)
        if total_pages == 0:
            raise EmptyDocumentError("pdf")

        pages_text: list[str] = []
        empty_pages = 0

        for i, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as e:
                logger.warning("PDF 第 %d 页提取失败: %s", i, e)
                page_text = ""

            stripped = page_text.strip()
            if stripped:
                pages_text.append(stripped)
            else:
                empty_pages += 1
                logger.debug("PDF 第 %d 页为空，已跳过", i)

        logger.info(
            "PDF 解析统计: total_pages=%d, non_empty=%d, empty_filtered=%d",
            total_pages, len(pages_text), empty_pages,
        )

        return "\n\n".join(pages_text)

    # ---- DOCX Parser ----

    def _parse_docx(self, file_bytes: bytes) -> str:
        """
        Extract text from DOCX using python-docx.

        Iterates over paragraphs and preserves paragraph breaks.
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ParserError("docx", "python-docx 库未安装，请执行: pip install python-docx")

        stream = io.BytesIO(file_bytes)
        try:
            doc = DocxDocument(stream)
        except Exception as e:
            raise ParserError("docx", f"无法读取 DOCX 文件: {e}")

        paragraphs: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                paragraphs.append(text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        if not paragraphs:
            raise EmptyDocumentError("docx")

        logger.info("DOCX 解析统计: paragraphs=%d", len(paragraphs))
        return "\n".join(paragraphs)

    # ---- XLSX Parser ----

    def _parse_xlsx(self, file_bytes: bytes) -> str:
        """
        Extract text from XLSX using openpyxl.

        Each worksheet is prefixed with a markdown-style heading
        ``## Sheet: {sheet_name}`` to preserve sheet context.
        Cell values within a row are joined with tab characters.
        Empty rows are skipped.
        """
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ParserError("xlsx", "openpyxl 库未安装，请执行: pip install openpyxl")

        stream = io.BytesIO(file_bytes)
        try:
            wb = load_workbook(stream, read_only=True, data_only=True)
        except Exception as e:
            raise ParserError("xlsx", f"无法读取 XLSX 文件: {e}")

        output_parts: list[str] = []
        total_sheets = len(wb.sheetnames)
        total_rows = 0

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Sheet heading
            output_parts.append(f"## Sheet: {sheet_name}")

            rows_in_sheet = 0
            for row in ws.iter_rows(values_only=True):
                # Extract non-None cell values, convert to string
                cell_values: list[str] = []
                for cell in row:
                    if cell is not None:
                        val = str(cell).strip()
                        if val:
                            cell_values.append(val)

                if cell_values:
                    output_parts.append("\t".join(cell_values))
                    rows_in_sheet += 1

            total_rows += rows_in_sheet
            logger.debug(
                "XLSX 工作表 '%s': rows=%d", sheet_name, rows_in_sheet
            )

        wb.close()

        if total_rows == 0:
            raise EmptyDocumentError("xlsx")

        logger.info(
            "XLSX 解析统计: sheets=%d, total_rows=%d",
            total_sheets, total_rows,
        )
        return "\n".join(output_parts)

    # ---- Plain Text (TXT / MD) ----

    def _parse_plain_text(self, file_bytes: bytes, file_type: str) -> str:
        """
        Extract text from plain-text files (TXT, MD).

        Attempts UTF-8 decoding first, then falls back to GBK
        (common for Chinese documents), then latin-1 as last resort.
        """
        # Try encodings in order of likelihood
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                text = file_bytes.decode(encoding)
                logger.debug(
                    "%s 文件使用 %s 编码解码成功", file_type.upper(), encoding
                )
                return text
            except UnicodeDecodeError:
                continue

        # This should never happen — latin-1 decodes anything
        raise ParserError(file_type, "无法识别文件编码，请确认文件未损坏")

    # ---- Post-processing ----

    def _postprocess(self, text: str) -> str:
        """
        Normalize extracted text.

        - Collapse 3+ consecutive newlines into 2.
        - Strip leading/trailing whitespace from the whole document.
        """
        import re

        # Collapse excessive blank lines (3+ → 2)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Strip overall whitespace
        text = text.strip()

        return text


# ============================================================
# Module-level convenience
# ============================================================

# Singleton instance
_parser: Optional[DocumentParser] = None


def get_parser() -> DocumentParser:
    """Return a cached DocumentParser singleton."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser


def parse_document(file_bytes: bytes, file_type: str) -> str:
    """Convenience function: parse document bytes and return text."""
    return get_parser().parse(file_bytes, file_type)
